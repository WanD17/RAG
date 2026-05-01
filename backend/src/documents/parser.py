"""Document parser — extracts text and tables from PDF/DOCX/TXT/MD files.

PDF pipeline (pdfplumber):
  - Crop top/bottom 6% of each page to remove running headers/footers
  - Extract plain text per page (preserving layout)
  - Extract tables separately as formatted text blocks

DOCX/TXT/MD: unchanged from previous implementation.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from loguru import logger


@dataclass
class ParsedDocument:
    page_texts: list[str]          # raw per-page text (before joining)
    text: str                       # full joined text
    tables: list[str] = field(default_factory=list)  # formatted table blocks


def _format_table(rows: list[list[str | None]], page_num: int, filename: str) -> str:
    """Convert a pdfplumber table (list of rows) to a readable text block."""
    if not rows:
        return ""

    cleaned = [[cell.strip() if cell else "" for cell in row] for row in rows]
    header, *body = cleaned

    # 2-column tables → key: value format
    if all(len(r) == 2 for r in cleaned):
        lines = [f"[Table — {filename}, page {page_num}]"]
        for row in cleaned:
            if any(row):
                lines.append(f"{row[0]}: {row[1]}")
        return "\n".join(lines)

    # Multi-column → pipe-separated
    lines = [f"[Table — {filename}, page {page_num}]"]
    lines.append(" | ".join(header))
    lines.append("-" * 40)
    for row in body:
        if any(row):
            lines.append(" | ".join(row))
    return "\n".join(lines)


def _crop_bbox(page) -> tuple[float, float, float, float]:
    """Return bbox cropping top/bottom 6% of the page."""
    h = page.height
    margin = h * 0.06
    return (0, margin, page.width, h - margin)


def parse_pdf(file_path: str) -> ParsedDocument:
    import pdfplumber

    filename = Path(file_path).name
    page_texts: list[str] = []
    tables: list[str] = []

    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            cropped = page.crop(_crop_bbox(page))

            # Extract tables first, then mask them out of text extraction
            raw_tables = cropped.extract_tables()
            for tbl in (raw_tables or []):
                if tbl and len(tbl) > 1:
                    formatted = _format_table(tbl, page_num, filename)
                    if formatted:
                        tables.append(formatted)

            text = cropped.extract_text(x_tolerance=2, y_tolerance=2) or ""
            if text.strip():
                page_texts.append(text.strip())

    full_text = "\n\n".join(page_texts)
    logger.debug(f"PDF parsed: {filename} — {len(page_texts)} pages, {len(tables)} tables")
    return ParsedDocument(page_texts=page_texts, text=full_text, tables=tables)


def parse_docx(file_path: str) -> ParsedDocument:
    from docx import Document

    doc = Document(file_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)
    return ParsedDocument(page_texts=[text], text=text)


def parse_text(file_path: str) -> ParsedDocument:
    with open(file_path, encoding="utf-8", errors="replace") as f:
        text = f.read()
    return ParsedDocument(page_texts=[text], text=text)


def parse_file(file_path: str, file_type: str) -> ParsedDocument:
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = file_type.lower().lstrip(".")
    logger.debug(f"Parsing file: {path.name} (type={ext})")

    try:
        if ext == "pdf":
            return parse_pdf(file_path)
        elif ext == "docx":
            return parse_docx(file_path)
        elif ext in ("txt", "md"):
            return parse_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    except Exception as e:
        logger.error(f"Failed to parse {path.name}: {e}")
        raise
